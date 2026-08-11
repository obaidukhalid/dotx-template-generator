"""
template_builder.py
-------------------
Converts a styling configuration dictionary into a real Word template (.dotx).

Everything is driven by the config. Edit the config, rebuild, get a new template.

Public entry point:
    build_template(config: dict, output_path: str) -> str
"""

import os
import re
import shutil
import tempfile
import zipfile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Mm, Pt, RGBColor, Twips

# ----------------------------------------------------------------------------
# Constants
# ----------------------------------------------------------------------------

PAGE_SIZES = {
    "A4": (Mm(210), Mm(297)),
    "Letter": (Emu(914400 * 8.5), Emu(914400 * 11)),
    "Legal": (Emu(914400 * 8.5), Emu(914400 * 14)),
    "A5": (Mm(148), Mm(210)),
}

DOCX_MAIN_CT = (
    "application/vnd.openxmlformats-officedocument."
    "wordprocessingml.document.main+xml"
)
DOTX_MAIN_CT = (
    "application/vnd.openxmlformats-officedocument."
    "wordprocessingml.template.main+xml"
)

HEADING_KEYS = ["heading_1", "heading_2", "heading_3", "heading_4"]


# ----------------------------------------------------------------------------
# Low level XML helpers
# ----------------------------------------------------------------------------

def _get_or_add(parent, tag):
    """Return the child element with `tag`, creating it if missing."""
    el = parent.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        parent.append(el)
    return el


def _clean_theme_attrs(element, attrs):
    """Strip theme attributes so explicit values are not overridden by the theme."""
    for attr in attrs:
        if element.get(qn(attr)) is not None:
            del element.attrib[qn(attr)]


def _force_font(style, font_name, size_pt=None, bold=None, italic=None,
                color_hex=None, all_caps=None):
    """
    Apply run formatting to a style and remove competing theme references.

    Built-in Word styles bind fonts and colours to the document theme. Setting
    style.font.name alone leaves the theme attribute in place and the theme wins.
    This strips those attributes first.
    """
    rpr = style.element.get_or_add_rPr()

    if font_name:
        rfonts = _get_or_add(rpr, "w:rFonts")
        _clean_theme_attrs(
            rfonts,
            ["w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"],
        )
        rfonts.set(qn("w:ascii"), font_name)
        rfonts.set(qn("w:hAnsi"), font_name)
        rfonts.set(qn("w:cs"), font_name)

    if color_hex:
        color = _get_or_add(rpr, "w:color")
        _clean_theme_attrs(
            color, ["w:themeColor", "w:themeShade", "w:themeTint"]
        )
        color.set(qn("w:val"), color_hex.upper().lstrip("#"))

    if size_pt is not None:
        style.font.size = Pt(float(size_pt))
    if bold is not None:
        style.font.bold = bool(bold)
    if italic is not None:
        style.font.italic = bool(italic)
    if all_caps is not None:
        style.font.all_caps = bool(all_caps)


def _apply_paragraph_format(style, space_before=None, space_after=None,
                            line_spacing=None, keep_with_next=None,
                            indent_left=None, alignment=None):
    """Apply paragraph level formatting. Spacing values are in twips."""
    pf = style.paragraph_format
    if space_before is not None:
        pf.space_before = Twips(int(space_before))
    if space_after is not None:
        pf.space_after = Twips(int(space_after))
    if line_spacing is not None:
        pf.line_spacing = float(line_spacing)
    if keep_with_next is not None:
        pf.keep_with_next = bool(keep_with_next)
    if indent_left is not None:
        pf.left_indent = Twips(int(indent_left))
    if alignment is not None:
        pf.alignment = alignment


def _ensure_style(doc, name, style_id=None, based_on="Normal"):
    """
    Return a paragraph style by name, creating it if Word has it only as latent.

    style_id matters for built-in styles such as TOC 1, because Word matches
    those by their style id rather than by display name.
    """
    from docx.enum.style import WD_STYLE_TYPE

    try:
        return doc.styles[name]
    except KeyError:
        pass

    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if style_id:
        style.element.set(qn("w:styleId"), style_id)
    try:
        style.base_style = doc.styles[based_on]
    except KeyError:
        pass
    return style


# ----------------------------------------------------------------------------
# Numbering (bullet glyphs and indents)
# ----------------------------------------------------------------------------

def _numbering_part(doc):
    """Return the numbering part element, or None if the package has none."""
    try:
        return doc.part.numbering_part.element
    except (AttributeError, KeyError, NotImplementedError):
        return None


def _abstract_num_for(numbering_el, num_id):
    """Resolve a w:numId to its w:abstractNum element."""
    for num in numbering_el.findall(qn("w:num")):
        if num.get(qn("w:numId")) == str(num_id):
            ref = num.find(qn("w:abstractNumId"))
            if ref is None:
                return None
            abstract_id = ref.get(qn("w:val"))
            for abstract in numbering_el.findall(qn("w:abstractNum")):
                if abstract.get(qn("w:abstractNumId")) == abstract_id:
                    return abstract
    return None


def _style_num_id(doc, style_name):
    """Read the w:numId a paragraph style points at."""
    try:
        style = doc.styles[style_name]
    except KeyError:
        return None
    ppr = style.element.find(qn("w:pPr"))
    if ppr is None:
        return None
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        return None
    num_id = num_pr.find(qn("w:numId"))
    return num_id.get(qn("w:val")) if num_id is not None else None


def _set_bullet(doc, style_name, bullet_char, font_name, indent_twips):
    """Rewrite the bullet glyph, font and indent for a bullet list style."""
    numbering_el = _numbering_part(doc)
    if numbering_el is None:
        return

    num_id = _style_num_id(doc, style_name)
    if num_id is None:
        return

    abstract = _abstract_num_for(numbering_el, num_id)
    if abstract is None:
        return

    lvl = abstract.find(qn("w:lvl"))
    if lvl is None:
        return

    num_fmt = _get_or_add(lvl, "w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")

    lvl_text = _get_or_add(lvl, "w:lvlText")
    lvl_text.set(qn("w:val"), bullet_char)

    # The glyph must come from a font that actually contains it.
    rpr = _get_or_add(lvl, "w:rPr")
    rfonts = _get_or_add(rpr, "w:rFonts")
    _clean_theme_attrs(
        rfonts, ["w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"]
    )
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:hint"), "default")

    ppr = _get_or_add(lvl, "w:pPr")
    ind = _get_or_add(ppr, "w:ind")
    left = int(indent_twips) + 360
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), "360")


# ----------------------------------------------------------------------------
# Heading numbering
# ----------------------------------------------------------------------------

HEADING_NUMBER_FORMATS = {
    "decimal": {
        "label": "1, 1.1, 1.1.1",
        "formats": ["decimal"] * 4,
        "texts": ["%1", "%1.%2", "%1.%2.%3", "%1.%2.%3.%4"],
    },
    "decimal_dot": {
        "label": "1., 1.1., 1.1.1.",
        "formats": ["decimal"] * 4,
        "texts": ["%1.", "%1.%2.", "%1.%2.%3.", "%1.%2.%3.%4."],
    },
    "chapter": {
        "label": "Chapter 1, 1.1, 1.1.1",
        "formats": ["decimal"] * 4,
        "texts": ["Chapter %1", "%1.%2", "%1.%2.%3", "%1.%2.%3.%4"],
    },
    "section": {
        "label": "Section 1, 1.1, 1.1.1",
        "formats": ["decimal"] * 4,
        "texts": ["Section %1", "%1.%2", "%1.%2.%3", "%1.%2.%3.%4"],
    },
    "outline": {
        "label": "I., A., 1., a.",
        "formats": ["upperRoman", "upperLetter", "decimal", "lowerLetter"],
        "texts": ["%1.", "%2.", "%3.", "%4."],
    },
    "legal": {
        "label": "1.0, 1.1, 1.1.1",
        "formats": ["decimal"] * 4,
        "texts": ["%1.0", "%1.%2", "%1.%2.%3", "%1.%2.%3.%4"],
    },
}


def _next_numbering_ids(numbering_el):
    """Return free abstractNumId and numId values."""
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering_el.findall(qn("w:abstractNum"))
        if el.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(el.get(qn("w:numId")))
        for el in numbering_el.findall(qn("w:num"))
        if el.get(qn("w:numId")) is not None
    ]
    return (
        max(abstract_ids) + 1 if abstract_ids else 0,
        max(num_ids) + 1 if num_ids else 1,
    )


def _build_heading_abstract_num(abstract_id, scheme, suffix, levels, indent):
    """Create the multilevel list definition that drives heading numbers."""
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))

    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), "0A1B2C3D")
    abstract.append(nsid)

    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)

    tmpl = OxmlElement("w:tmpl")
    tmpl.set(qn("w:val"), "0B1C2D3E")
    abstract.append(tmpl)

    name = OxmlElement("w:name")
    name.set(qn("w:val"), "HeadingNumbers")
    abstract.append(name)

    for index in range(4):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(index))

        # Child order follows the schema: start, numFmt, pStyle, suff,
        # lvlText, lvlJc, pPr.
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)

        numbered = index < levels
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), scheme["formats"][index] if numbered else "none")
        lvl.append(num_fmt)

        p_style = OxmlElement("w:pStyle")
        p_style.set(qn("w:val"), f"Heading{index + 1}")
        lvl.append(p_style)

        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), suffix)
        lvl.append(suff)

        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), scheme["texts"][index] if numbered else "")
        lvl.append(lvl_text)

        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        lvl.append(lvl_jc)

        ppr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        if indent:
            # Hanging indent that steps in with each level.
            left = 360 * (index + 1)
            ind.set(qn("w:left"), str(left))
            ind.set(qn("w:hanging"), "360")
        else:
            # Numbers and text both sit flush against the left margin.
            ind.set(qn("w:left"), "0")
            ind.set(qn("w:firstLine"), "0")
        ppr.append(ind)
        lvl.append(ppr)

        abstract.append(lvl)

    return abstract


def _apply_heading_numbering(doc, config):
    """
    Attach an automatic multilevel list to Heading 1 to 4.

    Word then numbers headings itself, keeps the numbers in sequence as
    sections are added, moved or deleted, and carries them into the contents
    list.
    """
    doc_cfg = config.get("document", {})
    if not doc_cfg.get("number_headings", False):
        return

    scheme = HEADING_NUMBER_FORMATS.get(doc_cfg.get("number_format", "decimal"))
    if scheme is None:
        return

    numbering_el = _numbering_part(doc)
    if numbering_el is None:
        return

    levels = int(doc_cfg.get("number_levels", 4))
    levels = max(1, min(4, levels))
    suffix = doc_cfg.get("number_suffix", "tab")
    if suffix not in ("tab", "space", "nothing"):
        suffix = "tab"
    indent = bool(doc_cfg.get("number_indent", False))

    abstract_id, num_id = _next_numbering_ids(numbering_el)
    abstract = _build_heading_abstract_num(
        abstract_id, scheme, suffix, levels, indent
    )

    # abstractNum elements must precede num elements in the part.
    first_num = numbering_el.find(qn("w:num"))
    if first_num is not None:
        first_num.addprevious(abstract)
    else:
        numbering_el.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    numbering_el.append(num)

    # Link each heading style to its level in the list.
    for index in range(levels):
        style = doc.styles[f"Heading {index + 1}"]
        ppr = style.element.get_or_add_pPr()
        num_pr = ppr.get_or_add_numPr()
        num_pr.get_or_add_ilvl().val = index
        num_pr.get_or_add_numId().val = num_id


# ----------------------------------------------------------------------------
# Table of contents field
# ----------------------------------------------------------------------------

def _add_toc_field(doc, levels="1-4", placeholder_text=None):
    """Insert a live TOC field. Word fills it in when the field is updated."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f'TOC \\o "{levels}" \\h \\z \\u'

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = placeholder_text or (
        "Right click here and choose Update Field to build the contents list."
    )

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    for element in (begin, instr, separate, text, end):
        run._r.append(element)

    return paragraph


def _enable_update_fields(doc):
    """Ask Word to refresh fields when the document is opened."""
    settings = doc.settings.element
    update = _get_or_add(settings, "w:updateFields")
    update.set(qn("w:val"), "true")


# ----------------------------------------------------------------------------
# Style construction
# ----------------------------------------------------------------------------

def _configure_styles(doc, config):
    """Write every configured style into the document style sheet."""
    styles = config["styles"]
    page = config["page_setup"]

    # Normal is the root that everything else inherits from.
    normal = doc.styles["Normal"]
    _force_font(
        normal,
        page.get("default_font", "Calibri"),
        size_pt=page.get("default_font_size", 11),
        color_hex=styles["paragraph"]["body"].get("color", "1E293B"),
    )
    _apply_paragraph_format(
        normal,
        space_before=0,
        space_after=styles["paragraph"]["body"].get("spacing_after", 80),
        line_spacing=styles["paragraph"]["body"].get("line_spacing", 1.15),
    )

    # Headings 1 to 4.
    for index, key in enumerate(HEADING_KEYS, start=1):
        spec = styles["headings"][key]
        style = doc.styles[f"Heading {index}"]
        _force_font(
            style,
            spec.get("font", "Calibri"),
            size_pt=spec.get("fontSize"),
            bold=spec.get("bold", True),
            italic=spec.get("italic", False),
            color_hex=spec.get("color"),
            all_caps=spec.get("all_caps", False),
        )
        _apply_paragraph_format(
            style,
            space_before=spec.get("spacing_before"),
            space_after=spec.get("spacing_after"),
            line_spacing=spec.get("line_spacing", 1.0),
            keep_with_next=True,
        )

    # Body text.
    body = styles["paragraph"]["body"]
    body_style = _ensure_style(doc, "Body Text", "BodyText")
    _force_font(
        body_style,
        body.get("font", "Calibri"),
        size_pt=body.get("fontSize"),
        bold=body.get("bold", False),
        italic=body.get("italic", False),
        color_hex=body.get("color"),
    )
    _apply_paragraph_format(
        body_style,
        space_before=body.get("spacing_before", 0),
        space_after=body.get("spacing_after", 80),
        line_spacing=body.get("line_spacing", 1.15),
    )

    # Caption.
    caption = styles["paragraph"]["caption"]
    caption_style = _ensure_style(doc, "Caption", "Caption")
    _force_font(
        caption_style,
        caption.get("font", "Calibri"),
        size_pt=caption.get("fontSize"),
        bold=caption.get("bold", False),
        italic=caption.get("italic", True),
        color_hex=caption.get("color"),
    )
    _apply_paragraph_format(
        caption_style,
        space_before=caption.get("spacing_before", 40),
        space_after=caption.get("spacing_after", 60),
        line_spacing=1.0,
    )

    # Quote block, optional in the config.
    quote = styles["paragraph"].get("quote")
    if quote:
        quote_style = _ensure_style(doc, "Quote", "Quote")
        _force_font(
            quote_style,
            quote.get("font", "Calibri"),
            size_pt=quote.get("fontSize"),
            bold=quote.get("bold", False),
            italic=quote.get("italic", True),
            color_hex=quote.get("color"),
        )
        _apply_paragraph_format(
            quote_style,
            space_before=quote.get("spacing_before", 120),
            space_after=quote.get("spacing_after", 120),
            line_spacing=quote.get("line_spacing", 1.15),
            indent_left=quote.get("indent", 720),
        )

    # Bullet levels.
    bullet_styles = {"level_1": "List Bullet", "level_2": "List Bullet 2"}
    for level_key, style_name in bullet_styles.items():
        spec = styles["bullets"].get(level_key)
        if not spec:
            continue
        style = _ensure_style(doc, style_name, style_name.replace(" ", ""))
        _force_font(
            style,
            spec.get("font", "Calibri"),
            size_pt=spec.get("fontSize"),
            color_hex=spec.get("color"),
        )
        _apply_paragraph_format(
            style,
            space_before=0,
            space_after=spec.get("spacing_after", 40),
            line_spacing=spec.get("line_spacing", 1.15),
        )
        _set_bullet(
            doc,
            style_name,
            spec.get("bullet_char", "\u2022"),
            spec.get("font", "Calibri"),
            spec.get("indent", 0),
        )

    # Contents list heading and entries.
    toc = styles["toc"]
    toc_heading = _ensure_style(doc, "TOC Heading", "TOCHeading")
    _force_font(
        toc_heading,
        toc.get("title_font", styles["headings"]["heading_1"].get("font")),
        size_pt=toc.get("title_fontSize"),
        bold=toc.get("title_bold", True),
        color_hex=toc.get("title_color"),
    )
    _apply_paragraph_format(
        toc_heading,
        space_before=0,
        space_after=toc.get("spacing_after", 100),
        keep_with_next=True,
    )

    for level in range(1, 5):
        entry_style = _ensure_style(doc, f"TOC {level}", f"TOC{level}")
        _force_font(
            entry_style,
            toc.get("entry_font", body.get("font", "Calibri")),
            size_pt=toc.get("entry_fontSize"),
            bold=(level == 1 and toc.get("entry_bold_level1", False)),
            color_hex=toc.get("entry_color"),
        )
        _apply_paragraph_format(
            entry_style,
            space_before=0,
            space_after=toc.get("entry_spacing_after", 40),
            line_spacing=1.0,
            indent_left=(level - 1) * 220,
        )


def _configure_page(doc, config):
    """Apply page size and margins."""
    page = config["page_setup"]
    section = doc.sections[0]

    width, height = PAGE_SIZES.get(page.get("page_size", "A4"), PAGE_SIZES["A4"])
    if str(page.get("orientation", "portrait")).lower() == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = height, width
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width, section.page_height = width, height

    margins = page.get("margins", {})
    section.top_margin = Twips(int(margins.get("top", 1440)))
    section.bottom_margin = Twips(int(margins.get("bottom", 1440)))
    section.left_margin = Twips(int(margins.get("left", 1440)))
    section.right_margin = Twips(int(margins.get("right", 1440)))


# ----------------------------------------------------------------------------
# Document body
# ----------------------------------------------------------------------------

def _add_cover(doc, config):
    """Build a cover page from the configured placeholder text."""
    doc_cfg = config.get("document", {})
    headings = config["styles"]["headings"]
    body = config["styles"]["paragraph"]["body"]

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Twips(1600)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(doc_cfg.get("title_placeholder", "Document Title"))
    run.font.size = Pt(float(headings["heading_1"].get("fontSize", 28)) + 4)
    run.font.bold = True
    run.font.name = headings["heading_1"].get("font", "Calibri")
    run.font.color.rgb = RGBColor.from_string(
        headings["heading_1"].get("color", "006600").upper()
    )
    title.paragraph_format.space_after = Twips(120)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(doc_cfg.get("subtitle_placeholder", "Subtitle"))
    run.font.size = Pt(float(headings["heading_2"].get("fontSize", 20)) - 4)
    run.font.name = headings["heading_2"].get("font", "Calibri")
    run.font.color.rgb = RGBColor.from_string(
        headings["heading_2"].get("color", "005200").upper()
    )
    subtitle.paragraph_format.space_after = Twips(900)

    for line in (
        doc_cfg.get("author_placeholder", "Author name"),
        doc_cfg.get("date_placeholder", "Date"),
        doc_cfg.get("reference_placeholder", "Document reference"),
    ):
        if not line:
            continue
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(line)
        run.font.size = Pt(float(body.get("fontSize", 11)))
        run.font.name = body.get("font", "Calibri")
        paragraph.paragraph_format.space_after = Twips(60)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _add_toc_section(doc, config):
    """Add the contents heading and the live field."""
    toc = config["styles"]["toc"]
    heading = doc.add_paragraph(toc.get("title", "Table of Contents"))
    heading.style = doc.styles["TOC Heading"]

    _add_toc_field(doc, levels=toc.get("levels", "1-4"))
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _add_style_guide(doc, config):
    """Add a worked example of every style so the template is self documenting."""
    bullets = config["styles"]["bullets"]
    has_level_2 = "level_2" in bullets

    def heading(text, level):
        paragraph = doc.add_paragraph(text)
        paragraph.style = doc.styles[f"Heading {level}"]
        return paragraph

    def body(text):
        paragraph = doc.add_paragraph(text)
        paragraph.style = doc.styles["Body Text"]
        return paragraph

    def bullet(text, level=1):
        name = "List Bullet" if level == 1 else "List Bullet 2"
        paragraph = doc.add_paragraph(text)
        paragraph.style = doc.styles[name]
        return paragraph

    heading("Style Guide", 1)
    body(
        "This section shows every style defined in this template. Delete it once "
        "you start writing your own content."
    )

    heading("Heading 2 example", 2)
    body(
        "Body Text is the default style for normal paragraphs. It carries the "
        "font, size, colour, line spacing and paragraph spacing set in the "
        "configuration."
    )

    heading("Heading 3 example", 3)
    body("Use Heading 3 for sub sections inside a Heading 2 block.")

    heading("Heading 4 example", 4)
    body("Use Heading 4 for the smallest named division.")

    heading("Lists", 2)
    body("Apply the List Bullet style to create bullet points:")
    bullet("First level bullet point")
    bullet("Another first level point")
    if has_level_2:
        bullet("Second level bullet point", 2)
        bullet("Another second level point", 2)
    bullet("Back to the first level")

    heading("Inline emphasis", 2)
    text_cfg = config["styles"].get("text", {})
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles["Body Text"]
    paragraph.add_run("Use ")
    strong = paragraph.add_run("bold text")
    strong.bold = True
    if text_cfg.get("strong", {}).get("color"):
        strong.font.color.rgb = RGBColor.from_string(
            text_cfg["strong"]["color"].upper()
        )
    paragraph.add_run(" for key terms, ")
    emphasis = paragraph.add_run("italic text")
    emphasis.italic = True
    paragraph.add_run(" for references, and ")
    highlight = paragraph.add_run("coloured text")
    highlight.bold = bool(text_cfg.get("highlight", {}).get("bold", False))
    if text_cfg.get("highlight", {}).get("color"):
        highlight.font.color.rgb = RGBColor.from_string(
            text_cfg["highlight"]["color"].upper()
        )
    paragraph.add_run(" to draw attention to an important point.")

    heading("Captions", 2)
    body("Place a caption directly under a figure or a table:")
    caption = doc.add_paragraph(
        "Figure 1: Captions use a smaller size and a muted colour."
    )
    caption.style = doc.styles["Caption"]

    if config["styles"]["paragraph"].get("quote"):
        heading("Quotes", 2)
        quote = doc.add_paragraph(
            "Quote style is indented and set apart from the surrounding text."
        )
        quote.style = doc.styles["Quote"]

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _add_outline(doc, config):
    """Add a starter section outline for the selected document type."""
    outlines = {
        "proposal": [
            "Executive Summary",
            "Background and Problem Statement",
            "Proposed Solution",
            "Scope of Work",
            "Timeline and Milestones",
            "Commercial Terms",
            "About Us",
        ],
        "report": [
            "Executive Summary",
            "Introduction",
            "Method",
            "Results",
            "Discussion",
            "Conclusions and Recommendations",
            "References",
        ],
        "guide": [
            "Overview",
            "System Requirements",
            "Installation",
            "Getting Started",
            "Feature Reference",
            "Troubleshooting",
            "Frequently Asked Questions",
        ],
        "study_guide": [
            "Learning Objectives",
            "Key Concepts",
            "Worked Examples",
            "Practice Questions",
            "Summary",
            "Glossary",
        ],
        "specification": [
            "Purpose and Scope",
            "Definitions and Abbreviations",
            "System Overview",
            "Functional Requirements",
            "Non Functional Requirements",
            "Interfaces",
            "Verification and Validation",
        ],
    }

    doc_type = config.get("document", {}).get("outline_type", "none")
    sections = outlines.get(doc_type)
    if not sections:
        return

    for title in sections:
        paragraph = doc.add_paragraph(title)
        paragraph.style = doc.styles["Heading 1"]
        placeholder = doc.add_paragraph("[Write this section.]")
        placeholder.style = doc.styles["Body Text"]


def _add_header_footer(doc, config):
    """Add optional header text and a page number footer."""
    doc_cfg = config.get("document", {})
    section = doc.sections[0]
    body = config["styles"]["paragraph"]["body"]
    muted = config["styles"]["paragraph"]["caption"].get("color", "6B7280")

    header_text = doc_cfg.get("header_text", "").strip()
    if header_text:
        paragraph = section.header.paragraphs[0]
        paragraph.text = header_text
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in paragraph.runs:
            run.font.size = Pt(9)
            run.font.name = body.get("font", "Calibri")
            run.font.color.rgb = RGBColor.from_string(muted.upper())

    if not doc_cfg.get("page_numbers", True):
        return

    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.font.size = Pt(9)
    run.font.name = body.get("font", "Calibri")
    run.font.color.rgb = RGBColor.from_string(muted.upper())

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, end):
        run._r.append(element)


# ----------------------------------------------------------------------------
# Package conversion
# ----------------------------------------------------------------------------

def _convert_to_dotx(path):
    """
    Rewrite the package content type so Word treats the file as a template.

    Renaming a .docx to .dotx is not enough. The main document part has to
    declare the template content type.
    """
    temp_fd, temp_path = tempfile.mkstemp(suffix=".zip")
    os.close(temp_fd)

    with zipfile.ZipFile(path, "r") as source:
        items = source.infolist()
        with zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as target:
            for item in items:
                data = source.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    text = data.decode("utf-8")
                    text = text.replace(DOCX_MAIN_CT, DOTX_MAIN_CT)
                    data = text.encode("utf-8")
                target.writestr(item, data)

    shutil.move(temp_path, path)
    return path


# ----------------------------------------------------------------------------
# Public entry point
# ----------------------------------------------------------------------------

def build_template(config, output_path):
    """
    Build a .dotx template from a config dictionary.

    Returns the absolute path of the file that was written.
    """
    output_path = os.path.abspath(output_path)
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

    doc = Document()

    _configure_page(doc, config)
    _configure_styles(doc, config)
    _apply_heading_numbering(doc, config)

    doc_cfg = config.get("document", {})
    if doc_cfg.get("include_cover", True):
        _add_cover(doc, config)
    if doc_cfg.get("include_toc", True):
        _add_toc_section(doc, config)
    if doc_cfg.get("include_style_guide", True):
        _add_style_guide(doc, config)
    _add_outline(doc, config)

    # A template with no body at all gives the user nowhere to start typing.
    if not doc.paragraphs:
        starter = doc.add_paragraph()
        starter.style = doc.styles["Body Text"]

    _add_header_footer(doc, config)

    _enable_update_fields(doc)

    meta = config.get("template", {})
    doc.core_properties.title = meta.get("name", "Document Template")
    doc.core_properties.comments = meta.get("description", "")
    doc.core_properties.author = meta.get("author", "")
    doc.core_properties.category = "Template"

    doc.save(output_path)

    if output_path.lower().endswith(".dotx"):
        _convert_to_dotx(output_path)

    return output_path
